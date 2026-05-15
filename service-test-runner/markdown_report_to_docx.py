#!/usr/bin/env python3
"""Convert the generated Chinese compatibility Markdown report to DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_md")
    parser.add_argument("output_docx")
    return parser.parse_args()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_width(table, col_widths_dxa: List[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(col_widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in col_widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_run(run, bold: bool = False, code: bool = False) -> None:
    run.bold = bold
    run.font.name = "Calibri" if not code else "Consolas"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), run.font.name)
    r_fonts.set(qn("w:hAnsi"), run.font.name)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5 if not code else 9.5)
    run.font.color.rgb = TEXT


def add_inline_markdown(paragraph, text: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            style_run(run)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            style_run(run, code=True)
        else:
            run = paragraph.add_run(token[2:-2])
            style_run(run, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        style_run(run)


def split_table_row(line: str) -> List[str]:
    cells = line.strip().strip("|").split("|")
    return [cell.strip().replace("---:", "").replace(":---", "").replace("---", "") for cell in cells]


def is_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(set(cell.replace(" ", "")) <= {"-", ":"} for cell in cells)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    if col_count == 5:
        widths = [3000, 1250, 1250, 1250, 2490]
    elif col_count == 4:
        widths = [4200, 1500, 1500, 2160]
    elif col_count == 2:
        widths = [3000, 6360]
    else:
        widths = [int(9360 / col_count)] * col_count

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            add_inline_markdown(paragraph, text)
            set_paragraph_spacing(paragraph, after=2, line=1.05)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
                for run in paragraph.runs:
                    run.bold = True
            if c_idx > 0 and re.match(r"^[0-9.]+%?$", text):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_fixed_table_width(table, widths)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Title", 22, RGBColor(11, 37, 69), 0, 10),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("MongoDB 兼容性测试报告")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    style_run(run)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    idx = 0
    pending_table: List[List[str]] = []

    def flush_table() -> None:
        nonlocal pending_table
        if pending_table:
            add_table(doc, pending_table)
            pending_table = []

    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if not stripped:
            flush_table()
            idx += 1
            continue
        if stripped.startswith("|"):
            if not is_separator(stripped):
                pending_table.append(split_table_row(stripped))
            idx += 1
            continue
        flush_table()
        if stripped.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            add_inline_markdown(paragraph, stripped[2:].strip())
        elif stripped.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline_markdown(paragraph, stripped[3:].strip())
        elif stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(paragraph, stripped[4:].strip())
        elif stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:].strip())
            set_paragraph_spacing(paragraph, after=4, line=1.167)
        else:
            paragraph = doc.add_paragraph()
            add_inline_markdown(paragraph, stripped)
            set_paragraph_spacing(paragraph)
        idx += 1
    flush_table()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


if __name__ == "__main__":
    args = parse_args()
    convert(Path(args.input_md), Path(args.output_docx))
